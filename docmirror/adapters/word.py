"""
Word Adapter — .docx → BaseResult
===================================

Extracts paragraphs and tables from Word documents using python-docx.

Processing logic:
    1. Opens the .docx file via ``python-docx.Document``.
    2. Iterates all paragraphs:
       - Paragraphs with a "Heading" style → ``title`` Block with heading_level
         inferred from the style name suffix (e.g., "Heading 2" → level=2).
       - All other non-empty paragraphs → ``text`` Block.
    3. Iterates all tables:
       - Each table row is read as a list of cell text values.
       - Empty rows (all cells blank) are skipped.
       - The resulting 2D list becomes a ``table`` Block.
    4. All Blocks are placed on a single PageLayout (page=0) since .docx
       does not have a reliable page-level structure at the parsing layer.

Metadata includes:
    - source_format: "docx"
    - paragraph_count: total paragraphs (including empty ones)
    - table_count: total tables in the document
"""

from __future__ import annotations

import logging
from pathlib import Path

from docmirror.framework.base import BaseParser, ParserOutput, ParserStatus
from docmirror.models.domain import BaseResult, Block, PageLayout

logger = logging.getLogger(__name__)


class WordAdapter(BaseParser):
    """Word (.docx) format adapter — extracts paragraphs and tables."""

    async def to_base_result(self, file_path: Path) -> BaseResult:
        """
        Parse a .docx file into a BaseResult.

        Uses python-docx to read document content. Heading paragraphs are
        converted to title Blocks with their heading level preserved;
        non-heading paragraphs become text Blocks; and tables become
        table Blocks with 2D list data.
        """
        from docx import Document
        doc = Document(str(file_path))

        blocks = []
        text_parts = []

        # Extract paragraphs — skip empty ones
        for para in doc.paragraphs:
            if not para.text.strip():
                continue

            # Detect heading style (e.g., "Heading 1", "Heading 2")
            btype = "title" if para.style and para.style.name.startswith("Heading") else "text"
            level = None
            if btype == "title":
                try:
                    level = int(para.style.name[-1])
                except (IndexError, ValueError):
                    level = 1

            blocks.append(Block(
                block_type=btype,
                raw_content=para.text,
                page=0,
                heading_level=level,
            ))
            text_parts.append(para.text)

        # Extract tables — each row becomes a list of cell text values
        for table in doc.tables:
            rows = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                if any(row_data):  # skip completely empty rows
                    rows.append(row_data)
            if rows:
                blocks.append(Block(block_type="table", raw_content=rows, page=0))

        page = PageLayout(page_number=0, blocks=tuple(blocks))
        metadata = {
            "source_format": "docx",
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
        }

        return BaseResult(pages=(page,), full_text="\n\n".join(text_parts), metadata=metadata)
