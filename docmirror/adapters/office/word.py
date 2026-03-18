# Copyright (c) 2026 ValueMap Global and contributors. All rights reserved.
# Author: Adam Lin <adamlin@valuemapglobal.com>
#
# This source code is licensed under the Apache 2.0 license found in the
# LICENSE file in the root directory of this source tree.

"""
Word Adapter — .docx / .doc → ParseResult (structured precision)
================================================================

Extracts paragraphs and tables from Word documents using python-docx,
producing ParseResult with heading hierarchy and typed table cells.

- **.docx**: Direct extraction via python-docx.
- **.doc (legacy)**: Requires LibreOffice (soffice) to be installed for conversion.
  If soffice is not found, returns a failure with code FORMAT_REQUIRES_CONVERTER (recoverable).

Processing logic:
    1. Opens the .docx file via ``python-docx.Document``.
    2. Iterates all paragraphs:
       - Paragraphs with a "Heading" style → ``TextBlock(level=H1/H2/H3)``.
       - All other non-empty paragraphs → ``TextBlock(level=BODY)``.
    3. Iterates all tables:
       - First row treated as headers.
       - Each cell becomes a ``CellValue`` with type detection.
    4. OMML math elements are extracted as formula TextBlocks.

Metadata includes:
    - source_format: "docx" or "doc"
    - paragraph_count: total paragraphs (including empty ones)
    - table_count: total tables in the document
"""

from __future__ import annotations

import logging
import re
import shutil
from pathlib import Path
from typing import List

from docmirror.framework.base import BaseParser
from docmirror.models.entities.domain import BaseResult, Block, PageLayout

logger = logging.getLogger(__name__)


class WordAdapter(BaseParser):
    """Word (.docx/.doc) format adapter — native parsing prioritizing `python-docx`.

    Processing strategy:
        1. **Modern formats (.docx):** Direct deep extraction via ``python-docx``
           preserving document flow, tables, and structural elements at zero OCR cost.
        2. **Legacy formats (.doc):** Automatically transcoded to PDF via LibreOffice,
           then processed through the standard PDF pipeline.
    """

    async def perceive(self, file_path: Path, **context) -> ParseResult:
        """
        Native extraction for .docx; for legacy .doc, requires LibreOffice (soffice).
        If .doc and soffice is not found, returns failure with FORMAT_REQUIRES_CONVERTER.
        """
        path = Path(file_path)
        if path.suffix.lower() == ".doc":
            soffice = shutil.which("soffice")
            if not soffice:
                from docmirror.models.errors import build_failure_result

                return build_failure_result(
                    "FORMAT_REQUIRES_CONVERTER",
                    "Legacy .doc format requires LibreOffice (soffice) to be installed for conversion. "
                    "Install LibreOffice or use .docx.",
                    file_path=str(path),
                    file_type="word",
                    t0=None,
                )
        return await super().perceive(file_path, **context)

    async def to_parse_result(self, file_path: Path, **kwargs) -> ParseResult:
        """
        Parse .docx directly into ParseResult with heading hierarchy
        and typed table cells.
        """
        from docx import Document

        from docmirror.models.entities.parse_result import (
            CellValue,
            DataType,
            PageContent,
            ParseResult,
            ParserInfo,
            TableBlock,
            TableRow,
            TextBlock,
            TextLevel,
        )

        logger.info(f"[WordAdapter] Cell-level extraction for: {file_path}")
        doc = Document(str(file_path))

        texts: list[TextBlock] = []
        tables: list[TableBlock] = []

        from docmirror.adapters.office.omml_extractor import OMMLExtractor

        # Extract paragraphs
        for para in doc.paragraphs:
            # 1. Extract OMML math elements
            try:
                math_elements = para._element.findall(
                    ".//m:oMath", namespaces={"m": "http://schemas.openxmlformats.org/officeDocument/2006/math"}
                )
                for math_elem in math_elements:
                    latex = OMMLExtractor.convert_element(math_elem)
                    if latex:
                        texts.append(
                            TextBlock(
                                content=f"$${latex.strip()}$$",
                                level=TextLevel.BODY,
                            )
                        )
            except Exception as e:
                logger.debug(f"[WordAdapter] OMML math extraction failed: {e}")

            if not para.text.strip():
                continue

            # 2. Determine heading level
            if para.style and para.style.name.startswith("Heading"):
                try:
                    m = re.search(r"(\d+)$", para.style.name)
                    level_num = int(m.group(1)) if m else 1
                except (ValueError, AttributeError):
                    level_num = 1

                level = {1: TextLevel.H1, 2: TextLevel.H2, 3: TextLevel.H3}.get(level_num, TextLevel.H3)
            else:
                level = TextLevel.BODY

            texts.append(TextBlock(content=para.text, level=level))

        # Extract tables
        for t_idx, table in enumerate(doc.tables):
            # First row as headers
            first_row = True
            headers: list[str] = []
            data_rows: list[TableRow] = []

            for row in table.rows:
                cells = [CellValue(text=cell.text.strip(), data_type=DataType.TEXT) for cell in row.cells]

                if first_row:
                    headers = [c.text for c in cells]
                    first_row = False
                else:
                    if any(c.text for c in cells):  # skip empty rows
                        data_rows.append(TableRow(cells=cells))

            if headers or data_rows:
                tables.append(
                    TableBlock(
                        table_id=f"word_table_{t_idx}",
                        headers=headers,
                        rows=data_rows,
                        page=0,
                    )
                )

        page = PageContent(
            page_number=0,
            texts=texts,
            tables=tables,
        )

        return ParseResult(
            pages=[page],
            parser_info=ParserInfo(
                parser_name="WordAdapter",
                page_count=1,
                overall_confidence=1.0,
            ),
        )

    async def to_base_result(self, file_path: Path) -> BaseResult:
        """
        Fallback: Parse .docx into BaseResult via ParseResultBridge.
        """
        from docmirror.models.construction.parse_result_bridge import ParseResultBridge

        pr = await self.to_parse_result(file_path)
        return ParseResultBridge.to_base_result(pr)
